"""
docx_inspector.py
从 .docx 文件中提取：
  - 纯文本
  - 批注（word/comments.xml）
  - 段落样式元数据（字体/字号/对齐）
  - format_summary：精简摘要字符串，供 AI 快速阅读
"""
from __future__ import annotations

import io
import zipfile
from typing import Union

from docx import Document
from docx.oxml.ns import qn


def inspect_docx(source: Union[str, bytes]) -> dict:
    """
    解析 .docx 文件，返回结构化信息。

    Args:
        source: 文件路径字符串 或 bytes 内容

    Returns:
        {
            "text": str,
            "comments": list[str],
            "styles_detected": list[dict],
            "format_summary": str,
        }
    """
    if isinstance(source, str):
        with open(source, 'rb') as f:
            raw = f.read()
    else:
        raw = source

    text             = _extract_text(raw)
    comments         = _extract_comments(raw)
    styles_detected  = _extract_styles(raw)
    format_summary   = _build_summary(styles_detected, comments)

    return {
        "text": text,
        "comments": comments,
        "styles_detected": styles_detected,
        "format_summary": format_summary,
    }


# ── 内部函数 ──────────────────────────────────────────

def _extract_text(raw: bytes) -> str:
    """提取纯文本（段落 + 表格）"""
    try:
        doc = Document(io.BytesIO(raw))
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_texts:
                    lines.append(" | ".join(row_texts))
        return "\n".join(lines)
    except Exception:
        return ""


def _extract_comments(raw: bytes) -> list:
    """从 word/comments.xml 提取批注文字"""
    comments = []
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if 'word/comments.xml' not in z.namelist():
                return comments
            xml_bytes = z.read('word/comments.xml')

        from lxml import etree
        root = etree.fromstring(xml_bytes)
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        for comment_el in root.findall(f'.//{{{ns}}}comment'):
            texts = ''.join(
                t.text for t in comment_el.findall(f'.//{{{ns}}}t')
                if t.text
            ).strip()
            if texts:
                comments.append(texts)
    except Exception:
        pass
    return comments


def _extract_styles(raw: bytes) -> list:
    """
    遍历文档段落，对每种 Word 样式名只采样一次，
    读取实际应用的字体/字号/对齐等属性。
    """
    styles_detected = []
    seen = set()
    try:
        doc = Document(io.BytesIO(raw))
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            style_name = (p.style.name if p.style else None) or 'Normal'
            if style_name in seen:
                continue
            seen.add(style_name)

            # 字号：优先 run 级别，再 style 级别
            font_size_pt = None
            for run in p.runs:
                if run.font.size:
                    font_size_pt = run.font.size.pt
                    break
            if font_size_pt is None and p.style and p.style.font.size:
                font_size_pt = p.style.font.size.pt

            # 中文字体（eastAsia）
            cn_font = None
            en_font = None
            if p.runs:
                rPr_el = p.runs[0]._element.find(qn('w:rPr'))
                if rPr_el is not None:
                    rFonts_el = rPr_el.find(qn('w:rFonts'))
                    if rFonts_el is not None:
                        cn_font = rFonts_el.get(qn('w:eastAsia'))
                        en_font = rFonts_el.get(qn('w:ascii')) or rFonts_el.get(qn('w:hAnsi'))
            # 回退到 style 字体
            if not en_font and p.style and p.style.font.name:
                en_font = p.style.font.name

            # 粗体
            bold = False
            if p.runs:
                bold = bool(p.runs[0].bold)
            if not bold and p.style and p.style.font.bold:
                bold = True

            # 对齐
            alignment = None
            if p.alignment is not None:
                alignment = str(p.alignment).split('.')[-1].lower()
            elif p.style and p.style.paragraph_format.alignment is not None:
                alignment = str(p.style.paragraph_format.alignment).split('.')[-1].lower()

            styles_detected.append({
                'style_name':   style_name,
                'sample_text':  p.text[:40],
                'font_size_pt': font_size_pt,
                'font_name_cn': cn_font,
                'font_name_en': en_font,
                'bold':         bold,
                'alignment':    alignment,
            })
    except Exception:
        pass
    return styles_detected


_ALIGN_CN = {
    'left':    '左对齐',
    'center':  '居中',
    'right':   '右对齐',
    'justify': '两端对齐',
}


def _build_summary(styles_detected: list, comments: list) -> str:
    """生成精简文字摘要，供 AI 快速提取关键格式信息"""
    lines = []

    if styles_detected:
        lines.append('[检测到的段落样式]')
        for s in styles_detected:
            parts = [s['style_name']]
            if s['font_size_pt']:
                parts.append(f"{s['font_size_pt']}pt")
            if s['font_name_cn']:
                parts.append(s['font_name_cn'])
            elif s['font_name_en']:
                parts.append(s['font_name_en'])
            if s['bold']:
                parts.append('粗体')
            if s['alignment']:
                parts.append(_ALIGN_CN.get(s['alignment'], s['alignment']))
            sample = s['sample_text']
            if sample:
                parts.append(f'（示例：{sample[:20]}）')
            lines.append('  ' + ' | '.join(parts))

    if comments:
        lines.append(f'[批注内容，共 {len(comments)} 条]')
        for i, c in enumerate(comments[:20], 1):
            lines.append(f'  {i}. {c}')

    return '\n'.join(lines)
